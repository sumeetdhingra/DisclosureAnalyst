"""Extract text/content from disclosure document files."""
from __future__ import annotations

import base64
import io
import mimetypes
import zipfile
from dataclasses import dataclass, field
from pathlib import Path


TEXT_EXTS = {".txt", ".md", ".csv", ".log", ".rtf", ".html", ".htm", ".xml", ".json"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
DOC_EXTS = {".doc"}
XLSX_EXTS = {".xlsx", ".xlsm"}
XLS_EXTS = {".xls"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".tif"}


@dataclass
class ExtractedFile:
    path: str
    kind: str  # "text", "image", "unreadable"
    text: str = ""
    image_b64: str = ""
    image_media_type: str = ""
    error: str = ""


@dataclass
class ExtractionResult:
    files: list[ExtractedFile] = field(default_factory=list)

    @property
    def text_files(self) -> list[ExtractedFile]:
        return [f for f in self.files if f.kind == "text"]

    @property
    def image_files(self) -> list[ExtractedFile]:
        return [f for f in self.files if f.kind == "image"]

    @property
    def unreadable_files(self) -> list[ExtractedFile]:
        return [f for f in self.files if f.kind == "unreadable"]


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for i, page in enumerate(reader.pages, 1):
        try:
            t = page.extract_text() or ""
        except Exception as e:
            t = f"[page {i} extract error: {e}]"
        if t.strip():
            parts.append(f"--- Page {i} ---\n{t}")
    return "\n\n".join(parts)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"--- Sheet: {sheet} ---")
        for row in ws.iter_rows(values_only=True):
            row_vals = ["" if v is None else str(v) for v in row]
            if any(v.strip() for v in row_vals):
                parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def _extract_xls(data: bytes) -> str:
    import xlrd

    book = xlrd.open_workbook(file_contents=data)
    parts = []
    for sheet in book.sheets():
        parts.append(f"--- Sheet: {sheet.name} ---")
        for r in range(sheet.nrows):
            row_vals = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
            if any(v.strip() for v in row_vals):
                parts.append(" | ".join(row_vals))
    return "\n".join(parts)


def _extract_text(data: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_zip(zip_path: Path) -> ExtractionResult:
    result = ExtractionResult()
    try:
        zf = zipfile.ZipFile(zip_path, "r")
    except zipfile.BadZipFile as e:
        result.files.append(ExtractedFile(path=str(zip_path), kind="unreadable",
                                          error=f"Invalid ZIP: {e}"))
        return result

    with zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            name = info.filename
            # skip mac metadata
            if "__MACOSX" in name or Path(name).name.startswith("._"):
                continue
            ext = Path(name).suffix.lower()
            try:
                data = zf.read(info)
            except Exception as e:
                result.files.append(ExtractedFile(path=name, kind="unreadable",
                                                  error=f"Read failed: {e}"))
                continue

            try:
                if ext in PDF_EXTS:
                    text = _extract_pdf(data)
                    result.files.append(ExtractedFile(path=name, kind="text", text=text))
                elif ext in DOCX_EXTS:
                    text = _extract_docx(data)
                    result.files.append(ExtractedFile(path=name, kind="text", text=text))
                elif ext in XLSX_EXTS:
                    text = _extract_xlsx(data)
                    result.files.append(ExtractedFile(path=name, kind="text", text=text))
                elif ext in XLS_EXTS:
                    text = _extract_xls(data)
                    result.files.append(ExtractedFile(path=name, kind="text", text=text))
                elif ext in TEXT_EXTS:
                    text = _extract_text(data)
                    result.files.append(ExtractedFile(path=name, kind="text", text=text))
                elif ext in IMAGE_EXTS:
                    media_type = mimetypes.guess_type(name)[0] or "image/jpeg"
                    if media_type == "image/jpg":
                        media_type = "image/jpeg"
                    if media_type not in {"image/jpeg", "image/png", "image/gif", "image/webp"}:
                        # Claude vision supports these four; skip others as unreadable text
                        result.files.append(ExtractedFile(path=name, kind="unreadable",
                                                          error=f"Unsupported image type {media_type}"))
                        continue
                    b64 = base64.standard_b64encode(data).decode("ascii")
                    result.files.append(ExtractedFile(path=name, kind="image",
                                                      image_b64=b64,
                                                      image_media_type=media_type))
                elif ext in DOC_EXTS:
                    result.files.append(ExtractedFile(path=name, kind="unreadable",
                                                      error="Legacy .doc format not supported (convert to .docx)"))
                else:
                    # Try as text fallback
                    try:
                        text = _extract_text(data)
                        if text.strip():
                            result.files.append(ExtractedFile(path=name, kind="text", text=text))
                        else:
                            result.files.append(ExtractedFile(path=name, kind="unreadable",
                                                              error=f"Unknown file type: {ext}"))
                    except Exception as e:
                        result.files.append(ExtractedFile(path=name, kind="unreadable",
                                                          error=f"Unknown type, text decode failed: {e}"))
            except Exception as e:
                result.files.append(ExtractedFile(path=name, kind="unreadable",
                                                  error=f"Extraction failed: {e}"))
    return result
